from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_FOLDER = Path(__file__).parent
PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
FONT_NAME = "Arial"
BODY_SIZE = 9.5
NAVY = "19324D"
BLUE = "2F6FDB"
TEAL = "159A9C"
GREEN = "27896A"
PURPLE = "6B5CA5"
ORANGE = "D97706"
LIGHT_BLUE = "EAF1FC"
LIGHT_GRAY = "F3F5F7"
WHITE = "FFFFFF"
DARK = "243447"
MUTED = "607080"


def set_run(run, size=BODY_SIZE, bold=False, color=DARK, font=FONT_NAME):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_fill(cell, color):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        item = margins.find(qn("w:" + side))
        if item is None:
            item = OxmlElement("w:" + side)
            margins.append(item)
        item.set(qn("w:w"), str(value))
        item.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=TABLE_INDENT_DXA):
    table.autofit = False
    properties = table._tbl.tblPr
    width_element = properties.first_child_found_in("w:tblW")
    width_element.set(qn("w:w"), str(sum(widths)))
    width_element.set(qn("w:type"), "dxa")
    indent_element = properties.first_child_found_in("w:tblInd")
    if indent_element is None:
        indent_element = OxmlElement("w:tblInd")
        properties.append(indent_element)
    indent_element.set(qn("w:w"), str(indent))
    indent_element.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            cell._tc.get_or_add_tcPr().tcW.set(qn("w:w"), str(widths[index]))
            cell._tc.get_or_add_tcPr().tcW.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row):
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def prepare_document(font=FONT_NAME):
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    normal = document.styles["Normal"]
    normal.font.name = font
    normal._element.rPr.rFonts.set(qn("w:ascii"), font)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    normal.font.size = Pt(BODY_SIZE)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.08

    for name, size in (("Heading 1", 13), ("Heading 2", 11)):
        style = document.styles[name]
        style.font.name = font
        style._element.rPr.rFonts.set(qn("w:ascii"), font)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), font)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(NAVY)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.keep_with_next = True

    bullet = document.styles["List Bullet"]
    bullet.font.name = font
    bullet._element.rPr.rFonts.set(qn("w:ascii"), font)
    bullet._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    bullet.font.size = Pt(BODY_SIZE)
    bullet.paragraph_format.left_indent = Inches(0.25)
    bullet.paragraph_format.first_line_indent = Inches(-0.15)
    bullet.paragraph_format.space_after = Pt(2)
    bullet.paragraph_format.line_spacing = 1.05
    return document


def add_text(document, text, size=BODY_SIZE, bold=False, color=DARK, align=None, after=3, font=FONT_NAME):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    if align is not None:
        paragraph.alignment = align
    set_run(paragraph.add_run(text), size, bold, color, font)
    return paragraph


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    set_run(paragraph.add_run(text))
    return paragraph


def add_heading(document, text, accent=TEAL):
    paragraph = document.add_paragraph(style="Heading 1")
    paragraph.paragraph_format.space_before = Pt(5)
    run = paragraph.add_run(text.upper())
    set_run(run, 12.5, True, accent)
    return paragraph


def add_contact_header(document, name, title, contact, accent=TEAL, centered=False, font=FONT_NAME):
    alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    add_text(document, name.upper(), 25, True, NAVY, alignment, 1, font)
    add_text(document, title.upper(), 12, True, accent, alignment, 3, font)
    add_text(document, contact, 9.2, False, MUTED, alignment, 6, font)


def add_job(document, title, company, dates, bullets):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    set_run(paragraph.add_run(title + " - " + company), 10, True, DARK)
    set_run(paragraph.add_run("   " + dates), 9, False, MUTED)
    for item in bullets:
        add_bullet(document, item)


def save(document, file_name):
    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(footer.add_run("Fictional resume created for SmartATS DOCX testing"), 7.5, False, MUTED)
    document.save(OUTPUT_FOLDER / file_name)


def create_ananya():
    document = prepare_document()
    add_contact_header(
        document,
        "Ananya Sharma",
        "Data Analyst",
        "+91 91000 20001 | ananya.sharma@example.com | Bengaluru, India",
        TEAL,
    )
    add_heading(document, "Professional Summary", TEAL)
    add_text(document, "Data analyst with 3 years of experience using Python, SQL, pandas, Microsoft Excel and Power BI to clean data, build dashboards and explain business trends.")
    add_heading(document, "Technical Skills", TEAL)
    add_bullet(document, "Analysis: Python, pandas, NumPy, SQL, Statistics, scikit-learn")
    add_bullet(document, "Reporting: Microsoft Excel, Power BI, Tableau, Data Visualization")
    add_bullet(document, "Tools: Git, Jupyter Notebook, PostgreSQL")
    add_heading(document, "Professional Experience", TEAL)
    add_job(document, "Data Analyst", "InsightWorks Solutions", "Jan 2023 - Present", [
        "Cleaned sales and operations datasets with Python and pandas for weekly reporting.",
        "Wrote SQL queries and built Power BI and Microsoft Excel dashboards used by managers.",
        "Applied statistics to investigate trends and maintained analysis files with Git.",
    ])
    add_heading(document, "Projects", TEAL)
    add_text(document, "Retail Performance Dashboard - Combined PostgreSQL data with SQL, analysed it in Python and pandas, and created Power BI visualizations. Documented data definitions in Git and added a basic scikit-learn sales forecast.")
    add_heading(document, "Education", TEAL)
    add_text(document, "B.Sc. Statistics - Bangalore University, 2022", 9.8, True)
    save(document, "01_ananya_sharma_classic_strong.docx")


def add_cell_line(cell, text, size=BODY_SIZE, bold=False, color=DARK, after=2):
    paragraph = cell.add_paragraph() if cell.paragraphs[0].text else cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(after)
    set_run(paragraph.add_run(text), size, bold, color)


def create_vikram():
    document = prepare_document()
    add_contact_header(
        document,
        "Vikram Patel",
        "Business Intelligence Analyst",
        "Phone: +91 91000 20002 | vikram.patel@example.com | Ahmedabad, India",
        BLUE,
    )
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2500, 6860])
    left, right = table.rows[0].cells
    set_cell_fill(left, NAVY)
    set_cell_fill(right, WHITE)
    add_cell_line(left, "SKILLS", 12, True, "35D0D2", 5)
    for text_value in ["Python", "SQL", "Microsoft Excel", "Power BI", "pandas", "Data Visualization", "Statistics", "Git", "Tableau"]:
        add_cell_line(left, text_value, 9.2, False, WHITE, 2)
    add_cell_line(left, "LANGUAGES", 12, True, "35D0D2", 5)
    add_cell_line(left, "English\nHindi\nGujarati", 9.2, False, WHITE, 2)

    add_cell_line(right, "PROFILE", 12, True, BLUE, 4)
    add_cell_line(right, "BI analyst with 2 years of experience preparing SQL reports and interactive Power BI dashboards. Uses Python, pandas and Excel for data cleaning and visualization.", 9.3, False, DARK, 5)
    add_cell_line(right, "WORK EXPERIENCE", 12, True, BLUE, 4)
    add_cell_line(right, "Junior BI Analyst - MetricLane Consulting", 10, True, DARK, 1)
    add_cell_line(right, "Jul 2024 - Present", 8.8, False, MUTED, 2)
    add_cell_line(right, "Created SQL queries, cleaned data with Python and pandas, and delivered Excel and Power BI reports. Used Git for version control and reviewed dashboard data quality with the team.", 9.3, False, DARK, 5)
    add_cell_line(right, "PROJECTS", 12, True, BLUE, 4)
    add_cell_line(right, "Customer Retention Dashboard", 10, True, DARK, 1)
    add_cell_line(right, "Analysed customer records using SQL, Python and pandas. Built Power BI charts showing churn by product and region, then documented the findings in Git.", 9.3, False, DARK, 5)
    add_cell_line(right, "EDUCATION", 12, True, BLUE, 4)
    add_cell_line(right, "B.Sc. Business Analytics - Gujarat University, 2024", 9.5, True, DARK, 2)
    save(document, "02_vikram_patel_sidebar_good.docx")


def create_sana():
    document = prepare_document()
    add_contact_header(
        document,
        "Sana Khan",
        "Junior Reporting Analyst",
        "+91 91000 20003 | sana.khan@example.com | Hyderabad, India",
        PURPLE,
        True,
    )
    add_heading(document, "Career Objective", PURPLE)
    add_text(document, "Junior analyst with one year of internship experience preparing reports with SQL, Microsoft Excel, Python and basic Power BI.")
    add_heading(document, "Technical Proficiency", PURPLE)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2600, 6760])
    headers = table.rows[0].cells
    for index, text_value in enumerate(("Category", "Knowledge")):
        set_cell_fill(headers[index], PURPLE)
        paragraph = headers[index].paragraphs[0]
        set_run(paragraph.add_run(text_value), 9.5, True, WHITE)
    set_repeat_table_header(table.rows[0])
    rows = [
        ("Programming", "Python, SQL, basic pandas"),
        ("Reporting", "Microsoft Excel, Power BI, Data Visualization"),
        ("Methods", "Data cleaning, descriptive Statistics"),
        ("Tools", "GitHub, Jupyter Notebook"),
    ]
    for label, knowledge in rows:
        cells = table.add_row().cells
        set_table_geometry(table, [2600, 6760])
        set_cell_fill(cells[0], LIGHT_GRAY)
        set_run(cells[0].paragraphs[0].add_run(label), 9.2, True, DARK)
        set_run(cells[1].paragraphs[0].add_run(knowledge), 9.2, False, DARK)
    add_heading(document, "Experience", PURPLE)
    add_job(document, "Data Reporting Intern", "DataSpring Academy", "Jun 2025 - Jun 2026", [
        "Prepared monthly Microsoft Excel reports and wrote SQL queries for training datasets.",
        "Used Python and pandas to remove duplicates and created basic Power BI charts.",
    ])
    add_heading(document, "Academic Projects", PURPLE)
    add_text(document, "Student Performance Report - Cleaned CSV data with Python and pandas, calculated descriptive statistics and presented trends in Microsoft Excel and Power BI.")
    add_heading(document, "Education", PURPLE)
    add_text(document, "B.Sc. Mathematics - Osmania University, 2025", 9.8, True)
    save(document, "03_sana_khan_skills_table_moderate.docx")


def create_rahul():
    document = prepare_document(font="Georgia")
    add_contact_header(
        document,
        "Rahul Das",
        "Reporting Assistant",
        "+91 91000 20004 | rahul.das@example.com | Kolkata, India",
        ORANGE,
        True,
        "Georgia",
    )
    add_text(document, "I am an early-career reporting assistant with about one year of internship and freelance experience. I mainly prepare Microsoft Excel workbooks, clean small CSV files and write simple SQL queries for regular office reports.", 10.3, False, DARK, None, 10, "Georgia")
    add_text(document, "During my internship at EastPoint Services from July 2025 to June 2026, I checked spreadsheet totals, prepared monthly summaries and created basic bar charts. I used SQL to filter customer records and completed the reports through manual office processes.", 10.3, False, DARK, None, 10, "Georgia")
    add_text(document, "For a personal expense report, I combined several CSV exports in Microsoft Excel, used formulas and pivot tables, and wrote a short explanation of spending trends. The work was completed individually with basic reporting methods.", 10.3, False, DARK, None, 10, "Georgia")
    add_text(document, "I completed a B.A. in Economics from Calcutta University in 2025. I am currently learning more advanced reporting and analytical methods through introductory online courses.", 10.3, False, DARK, None, 10, "Georgia")
    add_text(document, "My strengths are careful checking, documentation and clear communication. I speak English, Hindi and Bengali.", 10.3, False, DARK, None, 10, "Georgia")
    save(document, "04_rahul_das_paragraph_partial.docx")


def create_priya():
    document = prepare_document()
    add_contact_header(
        document,
        "Priya Nair",
        "Digital Marketing Coordinator",
        "+91 91000 20005 | priya.nair@example.com | Kochi, India",
        GREEN,
    )
    cards = document.add_table(rows=2, cols=2)
    cards.style = "Table Grid"
    set_table_geometry(cards, [4680, 4680])
    content = [
        ("PROFILE", "Marketing coordinator with 2 years of experience planning social media campaigns, writing content and coordinating brand activities."),
        ("CORE SKILLS", "Content writing, social media, Canva, campaign planning, communication and customer research."),
        ("EXPERIENCE", "Marketing Coordinator - Coastline Creative | Jan 2024 - Present\nPrepared campaign calendars and monthly engagement summaries. Collaborated with designers and account teams."),
        ("PROJECT", "Community Brand Launch\nPlanned posts, coordinated photography and presented campaign reach in a simple spreadsheet. The project focused entirely on marketing communication."),
    ]
    for cell, item in zip([cell for row in cards.rows for cell in row.cells], content):
        set_cell_fill(cell, LIGHT_GRAY)
        add_cell_line(cell, item[0], 11, True, GREEN, 4)
        add_cell_line(cell, item[1], 9.2, False, DARK, 3)
    add_heading(document, "Education", GREEN)
    add_text(document, "B.A. English Literature - Mahatma Gandhi University, 2023", 9.8, True)
    add_heading(document, "Achievements and Interests", GREEN)
    add_text(document, "Organized a college media festival | Volunteer photography | English and Malayalam")
    save(document, "05_priya_nair_modern_cards_low.docx")


def main():
    OUTPUT_FOLDER.mkdir(parents=True, exist_ok=True)
    create_ananya()
    create_vikram()
    create_sana()
    create_rahul()
    create_priya()
    print("Created five DOCX resumes in", OUTPUT_FOLDER)


if __name__ == "__main__":
    main()

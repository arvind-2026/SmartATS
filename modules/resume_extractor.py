from io import BytesIO
import logging

import pdfplumber
from docx import Document
from pypdf import PdfReader

import config
from modules.file_validator import get_file_extension


# Some valid PDFs have incomplete font metadata. Hide those library warnings
# while keeping errors and SmartATS extraction-quality checks visible.
logging.getLogger(config.PDF_FONT_LOGGER_NAME).setLevel(config.PDF_FONT_LOG_LEVEL)


def extract_txt_text(file_bytes):
    """Extract text from TXT bytes using the configured encodings."""

    for encoding in config.TEXT_ENCODINGS:
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue

    return ""


def extract_docx_text(file_bytes):
    """Extract paragraphs and table cells from a DOCX resume."""

    document = Document(BytesIO(file_bytes))
    text_parts = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text.strip())

    for table in document.tables:
        for row in table.rows:
            cell_texts = []

            for cell in row.cells:
                if cell.text.strip():
                    cell_texts.append(cell.text.strip())

            if cell_texts:
                text_parts.append(" | ".join(cell_texts))

    return "\n".join(text_parts)


def find_column_split(page):
    """Find a likely empty vertical gutter and return its page position."""

    words = page.extract_words()
    number_of_words = len(words)

    if number_of_words < config.MINIMUM_WORDS_FOR_COLUMN_CHECK:
        return None

    best_position = None
    best_score = None

    for percent in range(
        config.MINIMUM_COLUMN_SPLIT_PERCENT,
        config.MAXIMUM_COLUMN_SPLIT_PERCENT + 1,
        config.COLUMN_SCAN_STEP_PERCENT,
    ):
        position = page.width * percent / config.PERCENT_DIVISOR
        gutter_width = (
            page.width
            * config.GUTTER_CHECK_WIDTH_PERCENT
            / config.PERCENT_DIVISOR
        )
        left_words = 0
        right_words = 0
        gutter_words = 0
        crossing_words = 0

        for word in words:
            word_center = (word["x0"] + word["x1"]) / 2

            if word_center < position:
                left_words += 1
            else:
                right_words += 1

            if abs(word_center - position) <= gutter_width:
                gutter_words += 1

            if word["x0"] < position < word["x1"]:
                crossing_words += 1

        left_percent = left_words * config.PERCENT_DIVISOR / number_of_words
        right_percent = right_words * config.PERCENT_DIVISOR / number_of_words

        if left_percent < config.MINIMUM_WORDS_PER_COLUMN_PERCENT:
            continue

        if right_percent < config.MINIMUM_WORDS_PER_COLUMN_PERCENT:
            continue

        gutter_percent = gutter_words * config.PERCENT_DIVISOR / number_of_words
        crossing_percent = crossing_words * config.PERCENT_DIVISOR / number_of_words

        if gutter_percent > config.MAXIMUM_GUTTER_WORD_PERCENT:
            continue

        if crossing_percent > config.MAXIMUM_CROSSING_WORD_PERCENT:
            continue

        score = gutter_words + crossing_words

        if best_score is None or score < best_score:
            best_score = score
            best_position = position

    return best_position


def extract_page_text(page):
    """Automatically read a PDF page as one or two columns."""

    split_position = find_column_split(page)

    if split_position is None:
        return page.extract_text(layout=True) or ""

    left_box = (0, 0, split_position, page.height)
    right_box = (split_position, 0, page.width, page.height)
    left_text = page.crop(left_box).extract_text() or ""
    right_text = page.crop(right_box).extract_text() or ""
    page_parts = []

    if left_text.strip():
        page_parts.append(left_text.strip())

    if right_text.strip():
        page_parts.append(right_text.strip())

    return "\n\n".join(page_parts)


def extract_pdf_text(file_bytes):
    """Extract each PDF page using automatic layout detection."""

    reader = PdfReader(BytesIO(file_bytes))

    if reader.is_encrypted:
        password_result = reader.decrypt("")

        if password_result == 0:
            raise PermissionError(config.PASSWORD_FILE_MESSAGE)

    text_parts = []

    with pdfplumber.open(BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = extract_page_text(page)

            if page_text and page_text.strip():
                text_parts.append(page_text.strip())

    extracted_text = "\n\n".join(text_parts)

    if extracted_text.strip():
        return extracted_text

    # Use pypdf as a final fallback when pdfplumber finds no text.
    for page in reader.pages:
        page_text = page.extract_text()

        if page_text and page_text.strip():
            text_parts.append(page_text.strip())

    return "\n\n".join(text_parts)


def check_extraction_quality(text, file_extension):
    """Return a quality label and warning after extraction."""

    clean_text = text.strip()

    if len(clean_text) == 0:
        if file_extension == "pdf":
            return "Failed", config.SCANNED_FILE_MESSAGE

        return "Failed", config.EMPTY_FILE_MESSAGE

    if len(clean_text) < config.MINIMUM_TEXT_LENGTH:
        return "Low", config.LOW_TEXT_MESSAGE

    return "High", ""


def extract_resume_text(file_name, file_bytes):
    """Select the correct extractor and return a standard result dictionary."""

    file_extension = get_file_extension(file_name)
    result = {
        "success": False,
        "file_name": file_name,
        "file_extension": file_extension,
        "text": "",
        "quality": "Failed",
        "message": "",
    }

    try:
        if file_extension == "pdf":
            text = extract_pdf_text(file_bytes)
        elif file_extension == "docx":
            text = extract_docx_text(file_bytes)
        elif file_extension == "txt":
            text = extract_txt_text(file_bytes)
        else:
            result["message"] = config.INVALID_FILE_MESSAGE
            return result

        quality, warning = check_extraction_quality(text, file_extension)
        result["text"] = text.strip()
        result["quality"] = quality

        if quality == "Failed":
            result["message"] = warning
            return result

        result["success"] = True

        if warning:
            result["message"] = warning
        else:
            result["message"] = config.EXTRACTION_SUCCESS_MESSAGE

        return result


    except PermissionError as error:
        result["message"] = str(error)
        return result
    except Exception:
        result["message"] = config.CORRUPT_FILE_MESSAGE
        return result
